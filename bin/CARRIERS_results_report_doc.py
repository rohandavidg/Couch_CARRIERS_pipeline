#!/data5/bsi/epibreast/m087494.couch/Scripts/Progams/bcbio/anaconda/bin/python

"""
script to generate the CARRIERS bioinformatics 
report
"""

import argparse
import sys
import os
from docx import Document
from wand.image import Image
from wand.color import Color
from docx.shared import Inches
from bs4 import BeautifulSoup
import csv
from filter_cnv import parse_cnv_txt
from cruzdb import Genome
import subprocess

current_path = os.getcwd()

def main():
    args = parse_args()
    run(args.ggps_output, args.total_samples)
    

def run(ggps_output, total_samples):
    cnv_pdf, coverage_file, coverage_png, cnv_txt_dir, cnv_image_dir, run_name = ggps_paths(ggps_output)
    cnv_list = cnv_list_txt(cnv_txt_dir)
    create_cnv_filter_txt = generate_cnv_out(cnv_list)
    cnv_tmp_files = tmp_cnv_file(current_path, "_CNV.tmp.tsv")
    print cnv_tmp_files
    check_tmp_cnv(cnv_tmp_files)
    filtered_cnv_files = tmp_cnv_file(current_path, "filtered.cnv")
    cnv_file_image_dict = get_filtered_cnv_image(cnv_image_dir, filtered_cnv_files)
    create_heat_map = parse_pdf(cnv_pdf)
    parse_coverage_dict = find_failed_samples(coverage_file)
    get_line_count_dict = get_line_count(cnv_file_image_dict)
    write_basic_doc = create_word_document(coverage_png, run_name, 
                                           total_samples, parse_coverage_dict, cnv_file_image_dict, get_line_count_dict)


def parse_args():
    """
    parsing arguments
    """
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument('-g', dest='ggps_output', required=True,
                        help='full path of the GGPS output in delivery space')
    parser.add_argument('-t', dest='total_samples', required=True,
                        help='total number of samples in this GGPS run')
    args = parser.parse_args()
    return args


def ggps_paths(ggps_output):
    check_path(ggps_output)
    run_name = os.path.basename(ggps_output)
    cnv_pdf = ggps_output + "/" + "secondary/qc/cnv/Germline_QC_plot.pdf"
    coverage_file = ggps_output + "/" + "secondary/qc/ngsqc/SStats.flag.csv"
    coverage_png = ggps_output + "/" + "secondary/docs/coverage.png"
    cnv_txt_dir = ggps_output + "/" + "secondary/variants/cnv"
    cnv_image_dir = ggps_output + "/" + "secondary/qc/cnv"
    check_path(cnv_pdf)
    check_path(coverage_file)
    check_path(coverage_png)
    return (cnv_pdf, coverage_file, coverage_png, cnv_txt_dir, cnv_image_dir, run_name)
    

def check_path(some_path):
    if os.path.exists(some_path):
        pass
    else:
        print "{0} file not found, check GGPS output".format(some_path)
        sys.exit()
    
##################################FUNCTION TO FILTER CNV##################################################
def cnv_list_txt(cnv_dir):
    cnv_file_list = []
    for root, cnv_dir, files in os.walk(cnv_dir):
        for each_file in files:
            if each_file.endswith("CNV.txt") and each_file.startswith("s_"):
                cnv_file = root + "/" + each_file
                cnv_file_list.append(cnv_file)
    return cnv_file_list


def generate_cnv_out(cnv_file_list):
    for cnv_file in cnv_file_list:
        parse_cnv_txt(cnv_file)



def tmp_cnv_file(current_path, string):
    tmp_cnv_files = []
    for root, dirs, files in os.walk(current_path):
        for each_file in files:
            if each_file.endswith(string):
                tmp_cnv = root + "/" + each_file
                tmp_cnv_files.append(tmp_cnv)
    return tmp_cnv_files


def check_tmp_cnv(tmp_cnv_files):
    for cnv_file in tmp_cnv_files:
        print cnv_file
        sample_name = os.path.basename(cnv_file)
        sample_trunc = sample_name.split(".")[0]
        sample_name_dict =  parse_tmp_cnv_files(cnv_file)
        utr_check_dict  = check_cnv_utr(sample_trunc, sample_name_dict)
        write_out_cnv = write_out_utr_filtered_cnv(sample_trunc, utr_check_dict)
        print cnv_file
        os.remove(cnv_file)

def parse_tmp_cnv_files(cnv_file):
    csv_dict = {}
    with open(cnv_file, 'rb') as csvfile:
        header = csvfile.readline()
        new_header = header.strip().split("\t")
        reader = csv.DictReader(csvfile, fieldnames=new_header, delimiter="\t")
        for row in reader:
            gene = row['gene']
            key = row['chr'] + ":" + row['start.pos'] + "-" + row['stop.pos']
            csv_dict[key] = [gene, row['pval'], row['SNR.db'], row['CNV.log2ratio']]
    return csv_dict


def check_cnv_utr(sample_name, sample_dict):
    sample_update_dict = {}
    for key, value in sample_dict.items():
        chrom = key.split(":")[0]
        start_stop = key.split(":")[1]
        start = int(start_stop.split("-")[0])
        stop = int(start_stop.split("-")[1])
        gene = value[0]
        g = Genome(db="hg19")
        transcript = g.refGene.filter_by(name2=gene).first()
        try:
            UTR5 = [int(transcript.utr5[0]), int(transcript.utr5[1])]
            UTR3 = [int(transcript.utr3[0]), int(transcript.utr3[1])]
            if int(UTR5[0]) >= int(start) >= int(UTR5[1]) or int(UTR5[0]) >= int(stop) >= int(UTR5[1]) or int(UTR5[0]) <= int(start) <= int(UTR5[1]) or int(UTR5[0]) <= int(stop) <= int(UTR5[1]):
                pass
            elif int(UTR3[0]) >= int(start) >= int(UTR3[1]) or int(UTR3[0]) >= int(stop) >= int(UTR3[1]) or int(UTR3[0]) <= int(start) <= int(UTR3[1]) or int(UTR3[0]) <= int(stop) <= int(UTR3[1]):
                pass
            else:
                sample_update_dict[key] = value
        except (AttributeError, TypeError):
            print "cant find {0} gene for utr, as it's a possible psudo gene".format(gene) 
    return sample_update_dict


def write_out_utr_filtered_cnv(sample_name, sample_update_dict):
    csv_outfile = sample_name + ".filtered.cnv"
    if not sample_update_dict:
        pass
    else:
        with open(csv_outfile, 'wb') as fout:
            fout.write("chrom\tstart\tstop\tgene\tpval\tSNR.db\tCNV.log2ratio\n")
            for key, value in sample_update_dict.items():
                chrom = key.split(":")[0]
                start_stop = key.split(":")[1]
                start = int(start_stop.split("-")[0])
                stop = int(start_stop.split("-")[1])
                fout.write(chrom + "\t" + str(start) + "\t" + str(stop) + "\t" + "\t".join(str(i) for i in value) + "\n")



def get_filtered_cnv_image(cnv_image_dir, filtered_cnv_files):
    filtered_cnv_image_dict = {}
    for some_file in filtered_cnv_files:
        filtered_sample_name_file = os.path.basename(some_file)
        sample_name = filtered_sample_name_file.split(".")[0]
        for root, dirs, files in os.walk(cnv_image_dir):
            for some_image in files:
                if some_image.startswith(sample_name) and some_image.endswith("_CNV.png") and not some_image.endswith("_SNR_vs_CNV.png"):
                    CNV_image_file = root + "/" + some_image
                    filtered_cnv_image_dict[some_file] = CNV_image_file
    return filtered_cnv_image_dict


def parse_pdf(cnv_pdf):
    with Image(filename=cnv_pdf) as img:
        single_image = img.sequence[0]
        with Image(single_image) as i:
            i.resize(400, 500)
            i.save(filename="image.png")

####################################CHECKING FAILED SAMPLES FROM COVERAGE REPORT ##############################################

def find_failed_samples(coverage_csv):
    failed_samples_dict = {}
    with open(coverage_csv) as csv_fin:
        header = csv_fin.readline()
        new_header = header.strip().split(",")
        reader = csv.DictReader(csv_fin, fieldnames=new_header, delimiter=",")
        for row in reader:
            if float(row['TotalReads']) < 1 or float(row['TotalMappedReads']) < 1 or float(row['PercentInCaptureReads']) < float(0.5):
                failed_samples_dict[row['sampleID']] = [float(row['TotalReads']) * 1000000, float(row['TotalMappedReads']) * 1000000, 
                                                        float(row['ReadsInCaptureRegion']) * 1000000, float(row['PercentInCaptureReads']) * 100,
                                                        row['Indexes']]
    return failed_samples_dict


def get_line_count(filtered_cnv_image_dict):
    line_count_dict = {}
    for key, value in filtered_cnv_image_dict.items():
        with open(key) as fin:
            new_fin = fin.readlines()
            count = len(new_fin)
            line_count_dict[key] = count
    return line_count_dict


def create_word_document(coverage_png, run_name, total_samples, failed_samples_dict, filtered_cnv_image_dict, line_count_dict):
    document = Document()
    document.add_heading('BIOINFORMATICS_REPORT', 0)
    p = document.add_paragraph("RUN NAME:" + run_name)
    p = document.add_paragraph('SAMPLES SEQUENCED:' + str(total_samples))
    document.add_heading('FASTQC RESULTS', level=1)
    document.add_paragraph('The quality of reads after trimming 30 bases on the 5 prime end and 23 bases on the 3 prime ends looked good for all samples. The Q20 metric and Q30 metrics for the reads that were sequenced were found to be in an acceptable range')
    p = document.add_heading('COVERAGE RESULTS', level=1)
    if failed_samples_dict:
        total_failures = len(failed_samples_dict)
        document.add_paragraph('there are totally ' + str(total_failures) + " failures due to thresholds for coverage set to total reads less than 1 million, total reads in capture region less that 1 million, percent of reads in capture region less that 50%")
        for key, value in failed_samples_dict.items():
            document.add_paragraph(str(i) + '\t' for i  in value)
        document.add_picture(coverage_png, width=Inches(4.25))
    else:
        document.add_paragraph('There are no samples that failed due to coverage')
        document.add_picture(coverage_png, width=Inches(4.25))        
    p = document.add_heading('VARIANT ANALYSIS', level=1)
    document.add_paragraph('There were 4 Analysis plans used to report on the final set of variants. Each analysis plan is listed below')
    document.add_paragraph('1. ANALYSIS I')
    document.add_paragraph('The variant results were filtered based on the CARRIERS target region of interest and also any likely pathogenic with CAVA IMPACT set to HIGH and MODERATE. We also got rid of any common variants with noTCGA_ExAC_AF set to less than or equal to 0.03. There were a total of 98 variants after applying these filters')
    document.add_paragraph('   a. carriers target')
    document.add_paragraph('   b. CAVA IMPACT = HIGH and MODERATE')
    document.add_paragraph('   c. noTCGA_ExAc_INFO_AF <= 0.003')
    document.add_paragraph('\\ressrv06\PATH\COUCH\SEQUENCING\CARRIERS\Bioinformatics\RUNS\QIAGEN_TestRun_7_28_16\RE-RUN\Qiagen_CARRIERS.pilot_102_samples.RE-RUN.2016_Q2_Analysis1.xlsx')
    document.add_paragraph('2. ANALYSIS II')
    document.add_paragraph('The variant results were filtered based on the CARRIERS target region of interest and also any known pathogenic variant with Clinvar_20151008_GRCh37_pathogenic = 1. We also got rid of any common variants with noTCGA_ExAC_AF set to less than or equal to 0.03. There were a total of 19 variants after applying these filters')
    document.add_paragraph('   a. carriers target')
    document.add_paragraph('   b. Clinvar_20151008_GRCh37_pathogenic = 1')
    document.add_paragraph('   c. noTCGA_ExAc_INFO_AF <= 0.003')
    document.add_paragraph('\\ressrv06\PATH\COUCH\SEQUENCING\CARRIERS\Bioinformatics\RUNS\QIAGEN_TestRun_7_28_16\RE-RUN\Qiagen_CARRIERS.pilot_102_samples.RE-RUN.2016_Q2_Analysis2.xlsx')
    document.add_paragraph('3. ANALYSIS III')
    document.add_paragraph('The variant results were filtered based on the CARRIERS target region of interest and also any likely pathogenic with CAVA IMPACT set to HIGH and MODERATE')
    document.add_paragraph('   a. carriers target')
    document.add_paragraph('   b. CAVA IMPACT = HIGH and MODERATE')
    document.add_paragraph('\\ressrv06\PATH\COUCH\SEQUENCING\CARRIERS\Bioinformatics\RUNS\QIAGEN_TestRun_7_28_16\RE-RUN\Qiagen_CARRIERS.pilot_102_samples.RE-RUN.2016_Q2_Analysis3.xlsx')

###############WRITING CNV#########################################################################################
    p = document.add_heading('PATTERN CNV', level=1)
    document.add_picture("image.png")
    for key, value in filtered_cnv_image_dict.items():
        line_count = int(line_count_dict[key])
        document.add_picture(value, width=Inches(4.25))
        sample_name = os.path.basename(key)
        sample_truc = sample_name.split("(")[0]
        with open(key, 'rb') as csv_in:
#            line_count = sum(1 for _ in csv_in)
            header = csv_in.readline()
            new_header = header.strip().split("\t")
            reader = csv.DictReader(csv_in, fieldnames=new_header, delimiter="\t")
            table = document.add_table(rows=line_count, cols=7)
            table.autofit
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'chrom'
            hdr_cells[1].text = 'start'
            hdr_cells[2].text = 'stop'
            hdr_cells[3].text = 'gene'
            hdr_cells[4].text = 'pval'
            hdr_cells[5].text = 'SNR.db'
            hdr_cells[6].text = 'CNV.log2ratio'
            row_cells = table.add_row().cells
            for row in reader:
                if row['chrom']:
                    row_cells[0].text = str(row['chrom'])
                    row_cells[1].text = str(row['start'])
                    row_cells[2].text = str(row['stop'])
                    row_cells[3].text = str(row['gene'])
                    row_cells[4].text = str(row['pval'])
                    row_cells[5].text = str(row['SNR.db'])
                    row_cells[6].text = str(row['CNV.log2ratio'])
            document.add_page_break()
        os.remove(key)
    document.save('demo.docx')


if __name__ == "__main__":
    main()
